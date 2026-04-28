from datetime import date
from pathlib import Path
from textwrap import wrap

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parent
OUTPUT = ROOT / "Final_Btech_Major_Project_2026_DTU_Report.docx"
ASSET_DIR = ROOT / "outputs" / "sample_report_assets"
ASSET_DIR.mkdir(parents=True, exist_ok=True)

TITLE = (
    "Simulation-Based Implementation of Lean Manufacturing Techniques\n"
    "for Productivity Improvement in SMEs"
)
SUBTITLE = "A Data-Driven Case Study from a Beverage Bottling Line"
STUDENTS = ["Ashutosh Verma", "Divash Krishnam", "Raja"]
SUPERVISOR = "Dr. MOHD SHUAIB"
DEPARTMENT = "Production and Industrial Engineering"
SUPERVISOR_DEPT = "Professor"
UNIVERSITY = "DELHI TECHNOLOGICAL UNIVERSITY"
UNIVERSITY_SUB = "(Formerly Delhi College of Engineering)"
ADDRESS = "Bawana Road, Delhi-110042"
TODAY = date(2026, 4, 27)


def set_run_font(run, size=12, bold=False, italic=False):
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), "Times New Roman")
    rfonts.set(qn("w:hAnsi"), "Times New Roman")


def set_doc_defaults(doc):
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    for style_name, size in [("Title", 16), ("Heading 1", 14), ("Heading 2", 12), ("Heading 3", 12)]:
        style = doc.styles[style_name]
        style.font.name = "Times New Roman"
        style.font.size = Pt(size)
        style.font.bold = True
        rpr = style.element.get_or_add_rPr()
        rfonts = rpr.rFonts
        if rfonts is None:
            rfonts = OxmlElement("w:rFonts")
            rpr.append(rfonts)
        rfonts.set(qn("w:ascii"), "Times New Roman")
        rfonts.set(qn("w:hAnsi"), "Times New Roman")


def add_para(doc, text="", align=WD_ALIGN_PARAGRAPH.JUSTIFY, size=12, bold=False, italic=False, space_after=6):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.15
    if text:
        run = p.add_run(text)
        set_run_font(run, size=size, bold=bold, italic=italic)
    return p


def add_tab_line(doc, left, right, indent=0.0):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(indent)
    p.paragraph_format.tab_stops.add_tab_stop(Inches(6.0), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(f"{left}\t{right}")
    set_run_font(run, size=12)
    return p


def add_page_break(doc):
    doc.add_page_break()


def add_field_page_number(paragraph):
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")

    r1 = paragraph.add_run()._r
    r1.append(begin)
    r2 = paragraph.add_run()._r
    r2.append(instr)
    r3 = paragraph.add_run()._r
    r3.append(separate)
    r4 = paragraph.add_run()
    set_run_font(r4, size=11)
    r4._r.append(placeholder)
    r5 = paragraph.add_run()._r
    r5.append(end)


def set_page_number_format(section, fmt="decimal", start=1):
    sect_pr = section._sectPr
    pg_num = sect_pr.find(qn("w:pgNumType"))
    if pg_num is None:
        pg_num = OxmlElement("w:pgNumType")
        sect_pr.append(pg_num)
    pg_num.set(qn("w:start"), str(start))
    pg_num.set(qn("w:fmt"), fmt)


def add_footer_page_number(section, fmt="decimal", start=1):
    section.footer.is_linked_to_previous = False
    set_page_number_format(section, fmt=fmt, start=start)
    p = section.footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_field_page_number(p)


def add_university_header(doc, dept_line):
    add_para(doc, f"DEPARTMENT OF {dept_line.upper()}", align=WD_ALIGN_PARAGRAPH.CENTER, size=12, bold=True, space_after=0)
    add_para(doc, UNIVERSITY, align=WD_ALIGN_PARAGRAPH.CENTER, size=12, bold=True, space_after=0)
    add_para(doc, UNIVERSITY_SUB, align=WD_ALIGN_PARAGRAPH.CENTER, size=11, space_after=0)
    add_para(doc, ADDRESS, align=WD_ALIGN_PARAGRAPH.CENTER, size=11, space_after=12)


def add_chapter_heading(doc, number, title):
    add_para(doc, f"Chapter {number}", align=WD_ALIGN_PARAGRAPH.CENTER, size=16, bold=True, space_after=4)
    add_para(doc, title.upper(), align=WD_ALIGN_PARAGRAPH.CENTER, size=14, bold=True, space_after=16)


def add_subsection(doc, label, title):
    p = add_para(doc, "", align=WD_ALIGN_PARAGRAPH.LEFT, size=12, bold=True, space_after=4)
    run = p.add_run(f"{label} {title}")
    set_run_font(run, size=12, bold=True)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_run_font(run, size=12)
    return p


def add_numbered(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_run_font(run, size=12)
    return p


def load_font(name, size):
    candidates = {
        "regular": [Path("C:/Windows/Fonts/times.ttf"), Path("C:/Windows/Fonts/Times New Roman.ttf")],
        "bold": [Path("C:/Windows/Fonts/timesbd.ttf"), Path("C:/Windows/Fonts/times new roman bold.ttf")],
    }
    for candidate in candidates[name]:
        if candidate.exists():
            return ImageFont.truetype(str(candidate), size)
    return ImageFont.load_default()


def wrap_to_width(text, font, max_width, draw):
    words = str(text).split()
    if not words:
        return [""]
    lines = []
    current = words[0]
    for word in words[1:]:
        trial = f"{current} {word}"
        width = draw.textbbox((0, 0), trial, font=font)[2]
        if width <= max_width:
            current = trial
        else:
            lines.append(current)
            current = word
    lines.append(current)
    return lines


def make_table_image(name, headers, rows, rel_widths):
    regular = load_font("regular", 26)
    bold = load_font("bold", 28)
    margin = 40
    table_width = 1800
    usable = table_width - (2 * margin)
    widths = [int(usable * r) for r in rel_widths]
    widths[-1] += usable - sum(widths)

    scratch = Image.new("RGB", (table_width, 200), "white")
    draw = ImageDraw.Draw(scratch)

    row_heights = []
    cell_padding_x = 18
    cell_padding_y = 12

    header_height = 0
    wrapped_headers = []
    for idx, header in enumerate(headers):
        lines = wrap_to_width(header, bold, widths[idx] - 2 * cell_padding_x, draw)
        wrapped_headers.append(lines)
        header_height = max(header_height, len(lines) * 34 + 2 * cell_padding_y)
    row_heights.append(header_height)

    wrapped_rows = []
    for row in rows:
        wrapped = []
        row_height = 0
        for idx, value in enumerate(row):
            lines = wrap_to_width(value, regular, widths[idx] - 2 * cell_padding_x, draw)
            wrapped.append(lines)
            row_height = max(row_height, len(lines) * 31 + 2 * cell_padding_y)
        wrapped_rows.append(wrapped)
        row_heights.append(row_height)

    total_height = margin * 2 + sum(row_heights) + 4
    img = Image.new("RGB", (table_width, total_height), "white")
    draw = ImageDraw.Draw(img)

    x0 = margin
    y = margin

    # Header
    x = x0
    for idx, lines in enumerate(wrapped_headers):
        w = widths[idx]
        h = row_heights[0]
        draw.rectangle([x, y, x + w, y + h], outline="black", width=2, fill="#F3F3F3")
        text_y = y + cell_padding_y
        for line in lines:
            draw.text((x + cell_padding_x, text_y), line, font=bold, fill="black")
            text_y += 34
        x += w
    y += row_heights[0]

    # Body
    for row_index, wrapped in enumerate(wrapped_rows):
        x = x0
        h = row_heights[row_index + 1]
        for idx, lines in enumerate(wrapped):
            w = widths[idx]
            draw.rectangle([x, y, x + w, y + h], outline="black", width=2, fill="white")
            text_y = y + cell_padding_y
            for line in lines:
                draw.text((x + cell_padding_x, text_y), line, font=regular, fill="black")
                text_y += 31
            x += w
        y += h

    out = ASSET_DIR / name
    img.save(out)
    return out


def add_caption(doc, text, kind="figure"):
    align = WD_ALIGN_PARAGRAPH.CENTER
    italic = kind == "figure"
    p = add_para(doc, text, align=align, size=11, italic=italic, space_after=6)
    return p


def insert_image(doc, path, width=5.9):
    doc.add_picture(str(path), width=Inches(width))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER


def build():
    doc = Document()
    set_doc_defaults(doc)

    # Cover page
    for _ in range(2):
        add_para(doc, "", align=WD_ALIGN_PARAGRAPH.CENTER, size=12, space_after=10)
    add_para(doc, TITLE.split("\n")[0], align=WD_ALIGN_PARAGRAPH.CENTER, size=20, bold=True, space_after=2)
    add_para(doc, TITLE.split("\n")[1], align=WD_ALIGN_PARAGRAPH.CENTER, size=20, bold=True, space_after=18)
    add_para(doc, "A PROJECT REPORT", align=WD_ALIGN_PARAGRAPH.CENTER, size=14, bold=True, space_after=10)
    add_para(doc, "SUBMITTED IN PARTIAL FULFILLMENT OF THE REQUIREMENTS", align=WD_ALIGN_PARAGRAPH.CENTER, size=12, bold=True, space_after=2)
    add_para(doc, "FOR THE AWARD OF THE DEGREE", align=WD_ALIGN_PARAGRAPH.CENTER, size=12, bold=True, space_after=2)
    add_para(doc, "OF", align=WD_ALIGN_PARAGRAPH.CENTER, size=12, bold=True, space_after=2)
    add_para(doc, "BACHELOR OF TECHNOLOGY", align=WD_ALIGN_PARAGRAPH.CENTER, size=14, bold=True, space_after=2)
    add_para(doc, "IN", align=WD_ALIGN_PARAGRAPH.CENTER, size=12, bold=True, space_after=2)
    add_para(doc, DEPARTMENT.upper(), align=WD_ALIGN_PARAGRAPH.CENTER, size=13, bold=True, space_after=16)
    add_para(doc, "Submitted by:", align=WD_ALIGN_PARAGRAPH.CENTER, size=12, bold=True, space_after=4)
    for name in STUDENTS:
        add_para(doc, name, align=WD_ALIGN_PARAGRAPH.CENTER, size=12, space_after=2)
    add_para(doc, "Under the supervision of", align=WD_ALIGN_PARAGRAPH.CENTER, size=12, bold=True, space_after=4)
    add_para(doc, SUPERVISOR, align=WD_ALIGN_PARAGRAPH.CENTER, size=12, space_after=2)
    add_para(doc, f"DEPARTMENT OF {DEPARTMENT.upper()}", align=WD_ALIGN_PARAGRAPH.CENTER, size=12, bold=True, space_after=0)
    add_para(doc, UNIVERSITY, align=WD_ALIGN_PARAGRAPH.CENTER, size=12, bold=True, space_after=0)
    add_para(doc, UNIVERSITY_SUB, align=WD_ALIGN_PARAGRAPH.CENTER, size=11, space_after=0)
    add_para(doc, ADDRESS, align=WD_ALIGN_PARAGRAPH.CENTER, size=11, space_after=10)
    add_para(doc, TODAY.strftime("%B %Y").upper(), align=WD_ALIGN_PARAGRAPH.CENTER, size=12, bold=True, space_after=0)

    # Front matter section
    front = doc.add_section(WD_SECTION.NEW_PAGE)
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
    add_footer_page_number(front, fmt="lowerRoman", start=1)

    add_university_header(doc, DEPARTMENT)
    add_para(doc, "CANDIDATE'S DECLARATION", align=WD_ALIGN_PARAGRAPH.CENTER, size=14, bold=True, space_after=16)
    add_para(
        doc,
        "We, Divash Krishnam, Ashutosh Verma and Raja, students of Bachelor of Technology in Production and Industrial Engineering, hereby declare that the project report titled "
        f'"{TITLE.replace(chr(10), " ")}: {SUBTITLE}" '
        "submitted by us to the Department of Production and Industrial Engineering, Delhi Technological University, in partial fulfillment of the requirement for the award of the Bachelor of Technology degree, is based on our own academic study of the selected industrial dataset, simulation analysis and interpretation of results. This work has not previously formed the basis for the award of any degree, diploma, fellowship or similar recognition.",
    )
    add_para(doc, "Place: New Delhi", align=WD_ALIGN_PARAGRAPH.LEFT, size=12, space_after=2)
    add_para(doc, f"Date: {TODAY.strftime('%d/%m/%Y')}", align=WD_ALIGN_PARAGRAPH.LEFT, size=12, space_after=24)
    for name in STUDENTS:
        add_para(doc, name, align=WD_ALIGN_PARAGRAPH.RIGHT, size=12, space_after=2)

    add_page_break(doc)
    add_university_header(doc, DEPARTMENT)
    add_para(doc, "CERTIFICATE", align=WD_ALIGN_PARAGRAPH.CENTER, size=14, bold=True, space_after=16)
    add_para(
        doc,
        "I hereby certify that the project report titled "
        f'"{TITLE.replace(chr(10), " ")}: {SUBTITLE}" '
        "submitted by Divash Krishnam, Ashutosh Verma and Raja for fulfillment of the requirements for the award of the degree of Bachelor of Technology is a record of the project work carried out by the students under my guidance and supervision. To the best of my knowledge, this work has not been submitted in part or full for any degree or diploma to this University or elsewhere.",
    )
    add_para(doc, "Place: New Delhi", align=WD_ALIGN_PARAGRAPH.LEFT, size=12, space_after=2)
    add_para(doc, f"Date: {TODAY.strftime('%d/%m/%Y')}", align=WD_ALIGN_PARAGRAPH.LEFT, size=12, space_after=26)
    add_para(doc, SUPERVISOR, align=WD_ALIGN_PARAGRAPH.RIGHT, size=12, bold=True, space_after=2)
    add_para(doc, "(SUPERVISOR)", align=WD_ALIGN_PARAGRAPH.RIGHT, size=12, bold=True, space_after=2)
    add_para(doc, SUPERVISOR_DEPT, align=WD_ALIGN_PARAGRAPH.RIGHT, size=12, space_after=2)
    add_para(doc, f"Department of {DEPARTMENT}", align=WD_ALIGN_PARAGRAPH.RIGHT, size=12, space_after=2)
    add_para(doc, "Delhi Technological University", align=WD_ALIGN_PARAGRAPH.RIGHT, size=12, space_after=0)

    add_page_break(doc)
    add_para(doc, "ABSTRACT", align=WD_ALIGN_PARAGRAPH.CENTER, size=14, bold=True, space_after=16)
    add_para(doc, "Keywords - Lean manufacturing, productivity improvement, SME manufacturing, waiting waste, downtime analysis, total productive maintenance, 5S, bootstrap simulation.", align=WD_ALIGN_PARAGRAPH.LEFT, size=11, space_after=10)
    add_para(
        doc,
        "This report investigates how lean-manufacturing priorities can be selected in an SME beverage bottling line by using real production and downtime records instead of relying only on conceptual analysis. The dataset used in the study includes 59 production summaries, 265 hourly observations and 1,388 downtime events recorded between July 22, 2022 and February 6, 2023. To estimate the likely impact of lean interventions, the study applies a bootstrap-based simulation workflow and classifies downtime into micro-, minor- and major-stoppage groups.",
    )
    add_para(
        doc,
        "The observed line produced 212,358 liters, but only 42.66% of monitored time was value-adding production time. The remaining 57.64% was consumed by waiting waste. Micro-stoppages accounted for 61.31% of all events but only 13.59% of downtime minutes, whereas major stoppages represented only 6.77% of events and yet consumed 66.26% of lost time. Simulation results show that 5S with visual management improves output by 2.33%, TPM improves output by 20.61%, and an integrated lean bundle improves output by 32.19% while raising simulated availability from 42.37% to 55.33%.",
    )
    add_para(
        doc,
        "The main conclusion is that the most effective first action is to target the most costly stoppages rather than the most visible ones. In this case, TPM-focused intervention on major stoppages yields much greater productivity recovery than 5S alone.",
    )

    add_page_break(doc)
    add_para(doc, "ACKNOWLEDGEMENT", align=WD_ALIGN_PARAGRAPH.CENTER, size=14, bold=True, space_after=16)
    add_para(
        doc,
        f"We express our sincere gratitude to {SUPERVISOR}, {SUPERVISOR_DEPT}, for his valuable guidance, encouragement and continuous support during the course of this work. His suggestions helped us shape the problem statement, refine the analytical approach and improve the presentation of the results.",
    )
    add_para(
        doc,
        "We also acknowledge the support of the faculty members of the Department of Production and Industrial Engineering for providing an academic environment in which this project could be completed with clarity and discipline.",
    )
    add_para(
        doc,
        "Finally, we thank our families and well-wishers for their motivation and encouragement throughout the preparation of this project report.",
    )
    for name in STUDENTS:
        add_para(doc, name, align=WD_ALIGN_PARAGRAPH.RIGHT, size=12, space_after=2)

    add_page_break(doc)
    add_para(doc, "Contents", align=WD_ALIGN_PARAGRAPH.CENTER, size=14, bold=True, space_after=16)
    for left, right in [
        ("Candidate's Declaration", "i"),
        ("Certificate", "ii"),
        ("Abstract", "iii"),
        ("Acknowledgement", "iv"),
        ("List of Tables", "vi"),
        ("List of Figures", "vii"),
        ("List of Symbols, Abbreviations and Nomenclature", "viii"),
        ("CHAPTER 1: INTRODUCTION", "1"),
        ("CHAPTER 2: BACKGROUND AND LITERATURE REVIEW", "3"),
        ("CHAPTER 3: RESEARCH METHODOLOGY AND DATASET", "6"),
        ("CHAPTER 4: RESULTS AND ANALYSIS", "10"),
        ("CHAPTER 5: COMPARATIVE EVALUATION AND PRACTICAL IMPLICATIONS", "15"),
        ("CHAPTER 6: CONCLUSION", "20"),
        ("Appendix A", "22"),
        ("Appendix B", "24"),
        ("References", "25"),
    ]:
        add_tab_line(doc, left, right)

    add_page_break(doc)
    add_para(doc, "List of Tables", align=WD_ALIGN_PARAGRAPH.CENTER, size=14, bold=True, space_after=16)
    for left, right in [
        ("Table 2.1 : Literature Gap Summary", "5"),
        ("Table 3.1 : Dataset Characteristics", "4"),
        ("Table 4.1 : Baseline Performance Indicators", "10"),
        ("Table 4.2 : Downtime Category Structure", "12"),
        ("Table 4.3 : Simulated Scenario Outcomes", "14"),
        ("Table 4.4 : Simulation Uncertainty Range", "15"),
        ("Table 5.1 : Comparative Lean-Methodology Summary", "16"),
        ("Table 5.2 : Recommended Lean Implementation Roadmap", "18"),
        ("Table A.1 : Lean Scenario Assumptions", "22"),
        ("Table B.1 : Comparative Scoring Scale", "24"),
    ]:
        add_tab_line(doc, left, right)

    add_page_break(doc)
    add_para(doc, "List of Figures", align=WD_ALIGN_PARAGRAPH.CENTER, size=14, bold=True, space_after=16)
    for left, right in [
        ("Figure 4.1 : Monthly Average Operational Efficiency", "11"),
        ("Figure 4.2 : Downtime Event Duration Distribution", "11"),
        ("Figure 4.3 : Downtime Pareto by Event Category", "13"),
        ("Figure 4.4 : Pause Time versus Operational Efficiency", "13"),
        ("Figure 4.5 : Simulated Production Output by Scenario", "14"),
        ("Figure 4.6 : Waiting-Waste Reduction by Scenario", "14"),
        ("Figure 5.1 : KPI-Based Comparison of Lean-Evaluation Methodologies", "17"),
        ("Figure 5.2 : Industry-Fit Comparison of the Scored Methodologies", "18"),
        ("Figure A.1 : Production Rate Comparison by Product Size", "23"),
    ]:
        add_tab_line(doc, left, right)

    add_page_break(doc)
    add_para(doc, "LIST OF SYMBOLS, ABBREVIATIONS AND NOMENCLATURE", align=WD_ALIGN_PARAGRAPH.CENTER, size=14, bold=True, space_after=16)
    for symbol, meaning in [
        ("A", "Availability of the monitored production line"),
        ("T_op", "Operating time"),
        ("T_mon", "Monitored time"),
        ("T_dt", "Downtime / pause time"),
        ("Q", "Observed output quantity"),
        ("R", "Production rate during operating time"),
        ("SME", "Small and Medium Enterprise"),
        ("IIoT", "Industrial Internet of Things"),
        ("TPM", "Total Productive Maintenance"),
        ("5S", "Sort, Set in Order, Shine, Standardize and Sustain"),
    ]:
        add_tab_line(doc, symbol, meaning)

    # Main matter section
    main = doc.add_section(WD_SECTION.NEW_PAGE)
    add_footer_page_number(main, fmt="decimal", start=1)

    add_chapter_heading(doc, 1, "Introduction")
    add_subsection(doc, "1.1", "Overview")
    add_para(
        doc,
        "Productivity loss on a bottling line is not distributed evenly across operating problems. A plant may experience many short interruptions during a shift, but a much smaller number of long stoppages can consume most of the available production time. In lean terms, that lost time appears as waiting waste because monitored time passes without generating output.",
    )
    add_subsection(doc, "1.2", "Problem Formulation")
    add_para(
        doc,
        "For SMEs, choosing the first improvement action is difficult because maintenance capacity, manpower and budgets are limited. If the first intervention is selected poorly, effort may be spent without recovering much productive time. The problem addressed in this report is therefore to determine which lean intervention should be given priority on the observed beverage bottling line.",
    )
    add_subsection(doc, "1.3", "Objectives")
    for item in [
        "To analyze the production behavior and downtime structure of the observed bottling-line system.",
        "To identify the dominant waiting-waste pattern affecting productivity.",
        "To simulate the effect of selected lean interventions on throughput, availability and waste reduction.",
        "To recommend an implementation priority for lean techniques based on the observed loss pattern.",
    ]:
        add_bullet(doc, item)
    add_subsection(doc, "1.4", "Motivation")
    add_para(
        doc,
        "The study is motivated by the need to connect lean decisions to real plant data rather than to generic tool descriptions. Because the dataset contains both production output and event-level downtime records, it becomes possible to test which type of improvement recovers the most productive time in a reproducible way.",
    )

    add_page_break(doc)
    add_chapter_heading(doc, 2, "Background and Literature Review")
    add_subsection(doc, "2.1", "Lean Manufacturing and SME Context")
    add_para(
        doc,
        "Studies on lean manufacturing generally show positive links between lean practice and operational performance, but they also emphasize that SMEs face resource constraints, sequencing problems and measurement difficulties. This means lean tools should not be applied randomly; instead, they should be selected according to the most important losses in the plant.",
    )
    add_subsection(doc, "2.2", "Food and Beverage Production Context")
    add_para(
        doc,
        "Food and beverage lines are influenced by packaging synchronization, hygiene constraints, high line interdependence and frequent short interruptions. Research in this sector confirms the value of lean practice, but it rarely tells managers whether the first action should focus on micro-stoppages, major breakdowns or a combined improvement program.",
    )
    add_subsection(doc, "2.3", "Simulation and Downtime-Centered Studies")
    add_para(
        doc,
        "Simulation studies show that lean what-if analysis becomes stronger when process variability and system interaction are represented explicitly. Downtime-focused bottling studies likewise show that detailed event-level records reveal hidden inefficiencies. However, the literature still rarely combines SME relevance, beverage-line context, open event data and uncertainty-aware simulation in one workflow.",
    )
    add_subsection(doc, "2.4", "Research Gap")
    add_para(
        doc,
        "The key gap addressed in this report is the lack of a clear, reproducible method for converting observed stoppage patterns into a ranked lean-priority framework. This study addresses that gap by using real industrial data and bootstrap simulation to compare alternative lean actions directly.",
    )
    table_21 = make_table_image(
        "table_2_1.png",
        ["Theme", "Established Insight", "Gap Addressed in This Report"],
        [
            [
                "Lean in SMEs",
                "Lean practice improves operational performance, but implementation is constrained by limited resources and sequencing difficulties.",
                "Most studies stop at broad guidance and do not rank the first intervention from plant-level downtime evidence.",
            ],
            [
                "Food and beverage lines",
                "The sector is highly sensitive to interruptions, synchronization losses and hygiene-driven constraints.",
                "The literature rarely shows whether a bottling line should attack micro-stops or major stoppages first.",
            ],
            [
                "Simulation-based planning",
                "Simulation helps test lean scenarios before physical implementation and preserves process variability.",
                "Many published workflows use proprietary cases and are not easy for students or SMEs to reproduce.",
            ],
            [
                "Downtime-centered diagnosis",
                "Detailed event records reveal hidden inefficiencies and highlight severe stoppages that dominate lost time.",
                "Few studies convert stoppage structure into a directly actionable lean-priority sequence with uncertainty-aware output estimates.",
            ],
        ],
        [0.22, 0.34, 0.44],
    )
    add_caption(doc, "Table 2.1: Literature Gap Summary", kind="table")
    insert_image(doc, table_21, width=6.0)
    add_para(
        doc,
        "Viewed together, the literature suggests that lean tools are useful, but it does not provide a simple answer to the practical plant-level question faced by SME managers: which intervention should be implemented first when downtime losses are highly unequal. The present report responds to that question by linking event-level loss structure, scenario design and measurable productivity recovery in one DTU-style project report.",
    )

    add_page_break(doc)
    add_chapter_heading(doc, 3, "Research Methodology and Dataset")
    add_subsection(doc, "3.1", "Dataset Description")
    add_para(
        doc,
        "The analysis uses the Industrial Production Time-Series Dataset from a Beverage Bottling Line published on Zenodo. The dataset represents a real IIoT-monitored bottling operation and includes hourly production output records, daily production summaries, hourly operating breakdown records and individual downtime events. The monitoring period spans July 22, 2022 to February 6, 2023 and includes both 3-liter and 5-liter product sizes.",
    )
    add_subsection(doc, "3.2", "Analytical Framework")
    add_para(
        doc,
        "The study measures output, monitored time, operating time, pause time, availability and waiting-waste share. Downtime events are grouped into micro-stoppages (up to 2 minutes), minor stoppages (2 to 10 minutes) and major stoppages (above 10 minutes). These categories are then used to connect specific lean tools to specific loss structures.",
    )
    add_subsection(doc, "3.3", "Simulation Procedure")
    add_para(
        doc,
        "An event-level bootstrap simulation with 5,000 replications is used. Observed day profiles are resampled with replacement so that the variability seen in the real plant is preserved. Lean scenario rules are then applied to the sampled downtime events, and updated output, availability and waiting-waste levels are recalculated.",
    )
    add_subsection(doc, "3.4", "Scenario Design")
    add_para(
        doc,
        "Three scenarios are tested in the report: 5S with visual management, TPM and an integrated lean bundle. These scenario values are literature-grounded planning assumptions designed for comparison rather than guaranteed future outcomes.",
    )
    add_para(
        doc,
        "The 5S scenario assumes a 15% elimination of micro-stoppages, reflecting gains from workplace organization and faster response to routine disruptions. The TPM scenario assumes a 25% reduction in major-stoppage duration, representing stronger maintenance discipline and faster restoration after severe failures. The integrated bundle combines micro-, minor- and major-stop reductions to reflect a broader improvement program rather than a single isolated lean tool.",
    )

    add_page_break(doc)
    table_31 = make_table_image(
        "table_3_1.png",
        ["Attribute", "Value"],
        [
            ["Monitoring period", "July 22, 2022 to February 6, 2023"],
            ["Production summary records", "59"],
            ["Hourly operating observations", "265"],
            ["Downtime events", "1,388"],
            ["Observed product sizes", "3-liter and 5-liter containers"],
            ["Total recorded output", "212,358 liters from 60,888 units"],
            ["Monitored time", "238.65 hours"],
            ["Operating time", "101.81 hours"],
        ],
        [0.34, 0.66],
    )
    add_caption(doc, "Table 3.1: Dataset Characteristics", kind="table")
    insert_image(doc, table_31, width=6.0)
    add_para(
        doc,
        "The table above shows that the dataset is rich enough for plant-level analysis because it captures both production behavior and the detailed downtime events needed for event-based simulation.",
    )
    add_para(
        doc,
        "This combination of daily output records and event-level stoppage observations is what makes the analysis academically useful. It allows the report to explain not just that downtime exists, but how its internal structure changes the ranking of lean interventions.",
    )
    add_para(doc, "The key operational relationships used in the report are listed below:", bold=True)
    for item in [
        "Availability, A = T_op / T_mon",
        "Downtime, T_dt = T_mon - T_op",
        "Production Rate, R = Q / T_op",
        "Scenario Output, Q' = R x (T_mon - T'_dt)",
        "Waiting-Waste Reduction, DeltaW(%) = ((T_dt - T'_dt) / T_dt) x 100",
    ]:
        add_bullet(doc, item)

    add_page_break(doc)
    add_chapter_heading(doc, 4, "Results and Discussion")
    add_subsection(doc, "4.1", "Baseline Performance")
    add_para(
        doc,
        "The recorded line produced 212,358 liters from 60,888 units during 238.65 monitored hours. However, effective operating time was only 101.81 hours while pause time reached 137.56 hours. This means the line operated with 42.66% availability and lost 57.64% of monitored time as waiting waste.",
    )
    table_41 = make_table_image(
        "table_4_1.png",
        ["Metric", "Value"],
        [
            ["Total production", "212,358 liters"],
            ["Total units produced", "60,888"],
            ["Monitored time", "238.65 h"],
            ["Operating time", "101.81 h"],
            ["Pause time", "137.56 h"],
            ["Availability", "42.66%"],
            ["Waiting-waste share", "57.64%"],
            ["3-liter production rate", "1,963.71 L/h"],
            ["5-liter production rate", "2,918.68 L/h"],
        ],
        [0.42, 0.58],
    )
    add_caption(doc, "Table 4.1: Baseline Performance Indicators", kind="table")
    insert_image(doc, table_41, width=6.0)

    add_page_break(doc)
    insert_image(doc, ROOT / "outputs" / "monthly_efficiency.png", width=5.8)
    add_caption(doc, "Figure 4.1: Monthly average operational efficiency across the monitored period.", kind="figure")
    add_para(
        doc,
        "Monthly efficiency varies noticeably over the monitoring window. Lower efficiency appears in August 2022 and December 2022, while January 2023 shows the strongest operational performance. This confirms that the line does not behave like a stable deterministic system.",
    )
    insert_image(doc, ROOT / "outputs" / "downtime_distribution.png", width=5.8)
    add_caption(doc, "Figure 4.2: Downtime-event duration distribution.", kind="figure")

    add_page_break(doc)
    add_subsection(doc, "4.2", "Downtime Structure")
    table_42 = make_table_image(
        "table_4_2.png",
        ["Category", "Event count", "Share of downtime minutes"],
        [
            ["Micro (<= 2 min)", "851", "13.59%"],
            ["Minor (2 to 10 min)", "443", "20.15%"],
            ["Major (> 10 min)", "94", "66.26%"],
        ],
        [0.40, 0.22, 0.38],
    )
    add_caption(doc, "Table 4.2: Downtime Category Structure", kind="table")
    insert_image(doc, table_42, width=6.0)
    insert_image(doc, ROOT / "outputs" / "downtime_category_pareto.png", width=5.8)
    add_caption(doc, "Figure 4.3: Downtime Pareto by event category.", kind="figure")
    add_para(
        doc,
        "Micro-stoppages account for 61.31% of events but only 13.59% of downtime minutes, whereas major stoppages account for just 6.77% of events and yet consume 66.26% of lost time. The key insight is that event frequency alone does not identify the most damaging loss.",
    )
    insert_image(doc, ROOT / "outputs" / "pause_vs_efficiency.png", width=5.8)
    add_caption(doc, "Figure 4.4: Pause time versus operational efficiency.", kind="figure")

    add_page_break(doc)
    add_subsection(doc, "4.3", "Simulation Results")
    add_para(
        doc,
        "The simulation baseline produced 211,924.69 liters, which is very close to the actual observed output. This indicates that the bootstrap model preserves real plant behavior sufficiently well for scenario comparison.",
    )
    table_43 = make_table_image(
        "table_4_3.png",
        ["Scenario", "Output (L)", "Gain (%)", "Availability (%)", "Waiting waste recovered (h)"],
        [
            ["Baseline", "211,924.69", "0.00", "42.37", "0.00"],
            ["5S + Visual Management", "216,866.69", "2.33", "43.25", "2.10"],
            ["TPM", "255,604.10", "20.61", "50.81", "20.21"],
            ["Integrated Lean Bundle", "280,142.22", "32.19", "55.33", "30.99"],
        ],
        [0.30, 0.17, 0.14, 0.18, 0.21],
    )
    add_caption(doc, "Table 4.3: Simulated Scenario Outcomes", kind="table")
    insert_image(doc, table_43, width=6.0)
    insert_image(doc, ROOT / "outputs" / "scenario_output.png", width=5.8)
    add_caption(doc, "Figure 4.5: Simulated production output by scenario.", kind="figure")
    insert_image(doc, ROOT / "outputs" / "scenario_waste_reduction.png", width=5.8)
    add_caption(doc, "Figure 4.6: Waiting-waste reduction by scenario.", kind="figure")
    add_para(
        doc,
        "5S with visual management produces only a modest gain because it targets micro-stoppages, which are frequent but relatively low in lost-time severity. TPM yields a much larger productivity gain because it directly targets major stoppages. The integrated lean bundle performs best overall because it combines short-stop discipline with severe-stop reduction.",
    )
    table_44 = make_table_image(
        "table_4_4.png",
        ["Scenario", "P05 Output (L)", "P95 Output (L)", "Interpretation"],
        [
            ["Baseline", "188,572.81", "236,799.79", "Represents the central production range without lean intervention."],
            ["5S + Visual Management", "192,906.87", "242,100.46", "Shows modest but consistent improvement over the baseline."],
            ["TPM", "226,295.58", "287,389.24", "Reflects strong recovery when severe stoppages are reduced."],
            ["Integrated Lean Bundle", "247,828.10", "314,804.33", "Provides the strongest upper and lower performance range among all scenarios."],
        ],
        [0.25, 0.17, 0.17, 0.41],
    )
    add_caption(doc, "Table 4.4: Simulation Uncertainty Range", kind="table")
    insert_image(doc, table_44, width=6.0)
    add_para(
        doc,
        "The uncertainty range is important because it shows that the ranking of scenarios is not just a point-estimate effect. Even when sampling variation is taken into account, TPM and the integrated bundle remain materially stronger than 5S alone, which strengthens the credibility of the recommended priority order.",
    )
    add_subsection(doc, "4.4", "Implementation Recommendations")
    for item in [
        "First priority should be TPM, with emphasis on preventive maintenance, reliability tracking and rapid recovery from major stoppages.",
        "Second priority should be 5S and visual management to reduce routine micro-disruptions and improve workplace control.",
        "Third priority should be an integrated bundle after maintenance discipline is in place, so the plant can combine reliability gains with standardization benefits.",
    ]:
        add_bullet(doc, item)

    add_page_break(doc)
    add_chapter_heading(doc, 5, "Comparative Evaluation and Practical Implications")
    add_subsection(doc, "5.1", "Comparison with Other Methodologies")
    add_para(
        doc,
        "The proposed approach can be positioned against other lean-evaluation methodologies by comparing their KPI focus, strongest industry fit and practical accessibility for SME users. This comparison is interpretive rather than experimental, but it helps explain where the present report is especially strong and where its scope is narrower.",
    )
    table_51 = make_table_image(
        "table_5_1.png",
        ["Study", "Primary Focus", "Strongest Use Case", "Main Limitation"],
        [
            [
                "This study",
                "Throughput, availability, downtime and waiting-waste reduction",
                "Repetitive packaging and food-beverage SME lines with downtime logs",
                "Does not directly model quality, WIP, labor or cost, and scenario effects remain assumed rather than observed.",
            ],
            [
                "Possik et al. [4]",
                "Lean-tool sensitivity under uncertainty with a broad KPI set",
                "High-mix aerospace and context-sensitive systems",
                "Requires heavier co-simulation setup and is less accessible for low-resource SME users.",
            ],
            [
                "Oleghe and Salonitis [2]",
                "Human factors, policy feedback loops and simulation trade-offs",
                "Plants where workforce behavior and scheduling pressure shape outcomes",
                "Harder to calibrate and less directly tied to simple event-level downtime logs.",
            ],
            [
                "Reslan et al. [10]",
                "Digital-twin-enabled value stream analysis and process monitoring",
                "Digitally instrumented smart-factory cells",
                "Higher infrastructure cost and complexity than a typical SME can easily absorb.",
            ],
            [
                "Matindana and Shoshiwa [9]",
                "Sector-level lean adoption and waste reduction across SMEs",
                "Benchmarking food and beverage SMEs at cross-firm scale",
                "Not designed for plant-level what-if simulation or downtime-priority ranking.",
            ],
        ],
        [0.16, 0.22, 0.29, 0.33],
    )
    add_caption(doc, "Table 5.1: Comparative Lean-Methodology Summary", kind="table")
    insert_image(doc, table_51, width=6.0)
    insert_image(doc, ROOT / "outputs" / "methodology_kpi_heatmap.png", width=5.8)
    add_caption(doc, "Figure 5.1: KPI-based comparison of lean-evaluation methodologies used in the discussion.", kind="figure")
    add_para(
        doc,
        "The KPI heatmap shows that the present method is strongest when the decision problem is centered on throughput, downtime and waiting-waste recovery. It is intentionally narrower than broader operations-research frameworks that model labor balancing, defect formation or WIP dynamics, but it is also much easier to explain and reproduce in a student project environment.",
    )
    insert_image(doc, ROOT / "outputs" / "methodology_industry_fit_heatmap.png", width=5.8)
    add_caption(doc, "Figure 5.2: Industry-fit comparison of the scored methodologies.", kind="figure")
    add_para(
        doc,
        "The industry-fit comparison reinforces that the report's methodology is particularly well aligned with repetitive, machine-paced packaging systems such as beverage bottling, food packaging, dairy filling and similar FMCG-oriented lines. Its relative weakness appears in environments where the central problem is labor balancing, quality defects or live cyber-physical control.",
    )
    add_subsection(doc, "5.2", "Recommended Lean Implementation Roadmap")
    table_52 = make_table_image(
        "table_5_2.png",
        ["Priority", "Lean Focus", "Expected Impact", "Primary KPIs to Monitor"],
        [
            [
                "1",
                "TPM and maintenance discipline",
                "Largest reduction in severe stoppage minutes and the strongest recovery of waiting waste.",
                "Major-downtime minutes, waiting-waste hours, availability, MTTR, liters gained",
            ],
            [
                "2",
                "5S and visual management",
                "Reduction in routine micro-disruptions and better workplace control after major losses are stabilized.",
                "Micro-stop count, short-delay frequency, recovery time, pause reduction",
            ],
            [
                "3",
                "Integrated lean bundle",
                "Combined gains in reliability, stabilization and standardized work once the plant is ready for broader rollout.",
                "Output gain, availability gain, total pause time, waiting-waste share, OEE when quality data are added",
            ],
        ],
        [0.10, 0.24, 0.30, 0.36],
    )
    add_caption(doc, "Table 5.2: Recommended Lean Implementation Roadmap", kind="table")
    insert_image(doc, table_52, width=6.0)
    add_para(
        doc,
        "The roadmap above translates the simulation findings into a staged implementation sequence. It avoids the common mistake of treating all downtime losses as equally damaging and instead ties each lean step to the category of loss it is best suited to remove.",
    )
    add_subsection(doc, "5.3", "Limitations and Future Scope")
    add_para(
        doc,
        "The report has four important limitations. First, the lean improvements are scenario assumptions used for planning comparison rather than post-implementation observations from the same plant. Second, the dataset does not contain detailed cause-coded downtime labels, so mapping from downtime category to lean tool remains logically strong but operationally indirect. Third, the study focuses on one industrial case and therefore should not be generalized mechanically to all production systems. Fourth, the methodology-comparison scores are explanatory interpretation scores derived from published workflows rather than normalized experiments on one shared dataset.",
    )
    add_para(
        doc,
        "Future work can extend the framework by incorporating root-cause-coded downtime, changeover analysis, quality-loss information, workforce allocation, cost-benefit estimates and real post-intervention performance records. Those additions would make the recommendation engine even stronger for plant-level decision support.",
    )

    add_chapter_heading(doc, 6, "Conclusion")
    add_subsection(doc, "6.1", "Concluding Remarks")
    add_para(
        doc,
        "This project report shows that lean priorities in a beverage bottling line can be selected more effectively when they are linked to real production and downtime behavior rather than to generic tool descriptions. The dataset makes it possible to treat waiting waste as a measurable operating loss and to compare how much of that loss each lean scenario is capable of recovering.",
    )
    add_para(
        doc,
        "Across all comparisons, the main conclusion remains stable: micro-stoppages are the most frequent interruptions, but they are not the dominant source of lost time. Major stoppages consume the largest share of waiting waste, which is why TPM performs far better than 5S when the plant's first priority is productivity recovery.",
    )
    add_para(doc, "The simulation results can be summarized as follows:", bold=True)
    for item in [
        "5S plus visual management improves output by 2.33% and reduces waiting waste by 1.52%.",
        "TPM improves output by 20.61% and reduces waiting waste by 14.68%.",
        "The integrated lean bundle improves output by 32.19%, raises simulated availability from 42.37% to 55.33% and reduces waiting waste by 22.51%.",
    ]:
        add_bullet(doc, item)
    add_subsection(doc, "6.2", "Final Recommendation")
    add_para(
        doc,
        "For the observed bottling line, TPM should be the first lean intervention because it directly attacks the severe stoppages that dominate lost time. 5S and visual management remain valuable, but their strongest role is to support a more stable line after major downtime is brought under better control. In that sense, the report contributes a practical and reproducible way of converting observed plant losses into a defendable lean-priority sequence.",
    )

    add_page_break(doc)
    add_para(doc, "Appendix A", align=WD_ALIGN_PARAGRAPH.CENTER, size=14, bold=True, space_after=16)
    table_a1 = make_table_image(
        "table_a_1.png",
        ["Lean Scenario", "Assumption Used in Simulation", "Purpose"],
        [
            ["5S + Visual Management", "15% of micro-stoppages eliminated", "Represents better workplace organization and faster response to small disruptions."],
            ["TPM", "25% reduction in major stoppage duration", "Represents improved reliability and reduction of severe breakdown loss."],
            ["Integrated Lean Bundle", "20% micro-stop removal, 15% minor-stop reduction, 30% major-stop reduction", "Represents a combined program of workplace discipline and reliability improvement."],
        ],
        [0.24, 0.33, 0.43],
    )
    add_caption(doc, "Table A.1: Lean Scenario Assumptions", kind="table")
    insert_image(doc, table_a1, width=6.0)
    insert_image(doc, ROOT / "outputs" / "product_rate_comparison.png", width=5.5)
    add_caption(doc, "Figure A.1: Production rate comparison by product size.", kind="figure")
    add_para(
        doc,
        "The 5-liter product shows higher average and median production rates than the 3-liter product, indicating that product mix also influences line performance. This does not overturn the downtime-based lean ranking, but it suggests that future work could combine stoppage analysis with production scheduling and product-mix planning.",
    )

    add_page_break(doc)
    add_para(doc, "Appendix B", align=WD_ALIGN_PARAGRAPH.CENTER, size=14, bold=True, space_after=16)
    table_b1 = make_table_image(
        "table_b_1.png",
        ["Score", "Meaning"],
        [
            ["0", "Not explicitly modeled or measured in the cited methodology."],
            ["1", "Indirect, qualitative or weak emphasis."],
            ["2", "Explicit secondary emphasis."],
            ["3", "Core quantitative emphasis or strongest fit."],
        ],
        [0.16, 0.84],
    )
    add_caption(doc, "Table B.1: Comparative Scoring Scale", kind="table")
    insert_image(doc, table_b1, width=5.6)
    add_para(
        doc,
        "The methodology-comparison figures in Chapter 5 use the scoring scale shown above. The scores are author-assigned from the published methodology descriptions, abstracts and explicitly reported KPI scopes in the cited studies. They should therefore be read as comparative interpretation scores rather than normalized raw performance outcomes.",
    )
    add_para(
        doc,
        "Only papers that present a concrete lean-evaluation workflow were scored in the figures. Broad reviews and cross-sectional strategic studies informed the narrative review but were not plotted as plant-level methodology templates.",
    )

    add_page_break(doc)
    add_para(doc, "References", align=WD_ALIGN_PARAGRAPH.CENTER, size=14, bold=True, space_after=16)
    references = [
        "[1] I. Belekoukias, J. A. Garza-Reyes, and V. Kumar, The impact of lean methods and tools on the operational performance of manufacturing organisations, International Journal of Production Research, 2014.",
        "[2] O. Oleghe and K. Salonitis, Hybrid simulation modelling of the human-production process interface in lean manufacturing systems, International Journal of Lean Six Sigma, 2019.",
        "[3] F. Yu and Z. Chen, Tools, application areas and challenges of factory simulation in small and medium-sized enterprises: A review, Procedia CIRP, 2021.",
        "[4] J. Possik, A. Zouggar-Amrani, B. Vallespir, and G. Zacharewicz, Lean techniques impact evaluation methodology based on a co-simulation framework for manufacturing systems, International Journal of Computer Integrated Manufacturing, 2022.",
        "[5] S. Frecassetti, B. Kassem, K. Kundu, M. Ferrazzi, and A. Portioli-Staudacher, Introducing lean practices through simulation: A case study in an Italian SME, Quality Management Journal, 2023.",
        "[6] J. A. C. Bokhorst, W. Knol, J. Slomp, and T. Bortolotti, Assessing to what extent smart manufacturing builds on lean principles, International Journal of Production Economics, 2022.",
        "[7] G. A. David et al., IoT-Driven Deep Learning for Enhanced Industrial Production Forecasting, IEEE Internet of Things Journal, 2024.",
        "[8] G. A. David, P. M. C. Monson, and C. Soares Junior, Industrial Production Time-Series Dataset from a Beverage Bottling Line, Zenodo dataset, 2026.",
        "[9] J. M. Matindana and M. J. Shoshiwa, Lean manufacturing implementation in food and beverage SMEs in Tanzania: using structural equation modelling, Management System Engineering, 2025.",
        "[10] M. Reslan, M. J. Triebe, R. Venketesh, and A. J. Hartwell, Automation of Value Stream Mapping: A Case Study on Enhancing Lean Manufacturing Tools Through Digital Twins, Procedia CIRP, 2025.",
        "[11] C. Unal and S. Bilget, Examination of lean manufacturing systems by simulation technique in apparel industry, The Journal of the Textile Institute, 2021.",
        "[12] M. M. Shahriar, M. S. Parvez, M. A. Islam, and S. Talapatra, Implementation of 5S in a plastic bag manufacturing industry: A case study, Cleaner Engineering and Technology, 2022.",
        "[13] P. M. Rojasra and M. N. Qureshi, Performance Improvement through 5S in Small Scale Industry: A case study, International Journal of Modern Engineering Research, 2013.",
        "[14] C. M. L. Rahman, M. A. Hoque, and S. M. Uddin, Assessment of Total Productive Maintenance Implementation through Downtime and Mean Downtime Analysis, IOSR Journal of Engineering, 2014.",
        "[15] G. Pinto et al., TPM implementation and maintenance strategic plan: A case study, Procedia Manufacturing, 2020.",
        "[16] R. Shah and P. T. Ward, Lean manufacturing: context, practice bundles, and performance, Journal of Operations Management, 2003.",
        "[17] L. W. Friedman and H. H. Friedman, Analyzing Simulation Output Using the Bootstrap Method, Simulation, 1995.",
        "[18] A. Belhadi, Y. B. M. Sha'ri, F. E. Touriki, and S. El Fezazi, Lean production in SMEs: literature review and reflection on future challenges, Journal of Industrial and Production Engineering, 2018.",
        "[19] M. Dora, M. Kumar, D. Van Goubergen, A. Molnar, and X. Gellynck, Operational performance and critical success factors of lean manufacturing in European food processing SMEs, Trends in Food Science and Technology, 2013.",
        "[20] T. Al-Hawari, F. Aqlan, M. Al-Buhaisi, and Z. Al-Faqeer, Simulation-based analysis and productivity improvement of a fully automatic bottle-filling production system: A practical case study, 2010 International Conference on Computer Modeling and Simulation.",
        "[21] I. Zennaro, D. Battini, F. Sgarbossa, A. Persona, and R. De Marchi, Micro downtime: Data collection, analysis and impact on OEE in bottling lines, International Journal of Quality and Reliability Management, 2018.",
        "[22] G. K. Inyiama and S. A. Oke, Maintenance downtime evaluation in a process bottling plant, International Journal of Quality and Reliability Management, 2020.",
        "[23] C. Cuggia-Jimenez, E. Orozco-Acosta, and D. Mendoza-Galvis, Lean manufacturing: a systematic review in the food industry, Informacion Tecnologica, 2020.",
        "[24] J. P. Womack and D. T. Jones, Lean Thinking: Banish Waste and Create Wealth in Your Corporation, Simon and Schuster, 1996.",
        "[25] R. Shah and P. T. Ward, Defining and developing measures of lean production, Journal of Operations Management, 2007.",
        "[26] J. Bhamu and K. S. Sangwan, Lean manufacturing: literature review and research issues, International Journal of Operations and Production Management, 2014.",
        "[27] T. Melton, The benefits of lean manufacturing: What lean thinking has to offer the process industries, Chemical Engineering Research and Design, 2005.",
        "[28] S. Nakajima, Introduction to TPM: Total Productive Maintenance, Productivity Press, 1988.",
        "[29] P. Muchiri and L. Pintelon, Performance measurement using overall equipment effectiveness: literature review and practical application discussion, International Journal of Production Research, 2008.",
        "[30] B. Dal, P. Tugwell, and R. Greatbanks, Overall equipment effectiveness as a measure of operational improvement: A practical analysis, International Journal of Operations and Production Management, 2000.",
        "[31] L. del C. Ng Corrales, M. P. Lamban, M. E. Hernandez Korner, and J. Royo, Overall equipment effectiveness: systematic literature review and overview of different approaches, Applied Sciences, 2020.",
        "[32] I. P. S. Ahuja and J. S. Khamba, Total productive maintenance: literature review and directions, International Journal of Quality and Reliability Management, 2008.",
        "[33] F. A. Abdulmalek and J. Rajgopal, Analyzing the benefits of lean manufacturing and value stream mapping via simulation: A process sector case study, International Journal of Production Economics, 2007.",
    ]
    for ref in references:
        add_para(doc, ref, align=WD_ALIGN_PARAGRAPH.LEFT, size=11, space_after=4)

    doc.save(OUTPUT)


if __name__ == "__main__":
    build()
    print(f"Created: {OUTPUT}")
