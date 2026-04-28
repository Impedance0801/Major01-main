from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
OUTPUT = ROOT / "research_project_report.docx"
ACCENT = RGBColor(30, 72, 104)
SUBTLE = RGBColor(92, 117, 138)
BODY = RGBColor(30, 34, 38)


def set_run_font(run, name, size, bold=False, italic=False, color=BODY):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = color
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), name)
    rfonts.set(qn("w:hAnsi"), name)


def configure(doc):
    for section in doc.sections:
        section.top_margin = Inches(0.9)
        section.bottom_margin = Inches(0.85)
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(11.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    for style_name, font_name, size, color in [
        ("Title", "Cambria", 22, ACCENT),
        ("Heading 1", "Cambria", 15, ACCENT),
        ("Heading 2", "Cambria", 12.5, ACCENT),
    ]:
        style = doc.styles[style_name]
        style.font.name = font_name
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        rpr = style.element.get_or_add_rPr()
        rfonts = rpr.rFonts
        if rfonts is None:
            rfonts = OxmlElement("w:rFonts")
            rpr.append(rfonts)
        rfonts.set(qn("w:ascii"), font_name)
        rfonts.set(qn("w:hAnsi"), font_name)


def add_text(paragraph, text, font="Times New Roman", size=11.5, bold=False, italic=False, color=BODY):
    run = paragraph.add_run(text)
    set_run_font(run, font, size, bold=bold, italic=italic, color=color)
    return run


def centered(doc, text, size=12, bold=False, color=BODY, font="Cambria", space_after=4):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(space_after)
    add_text(p, text, font=font, size=size, bold=bold, color=color)
    return p


def body(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(7)
    p.paragraph_format.line_spacing = 1.15
    add_text(p, text)
    return p


def heading(doc, text):
    p = doc.add_paragraph()
    p.style = "Heading 1"
    add_text(p, text, font="Cambria", size=15, bold=True, color=ACCENT)
    return p


def subheading(doc, text):
    p = doc.add_paragraph()
    p.style = "Heading 2"
    add_text(p, text, font="Cambria", size=12.5, bold=True, color=ACCENT)
    return p


def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    add_text(p, text)
    return p


def number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    add_text(p, text)
    return p


def key_point(doc, label, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(5)
    add_text(p, f"{label}: ", font="Cambria", size=11, bold=True, color=ACCENT)
    add_text(p, text, size=11)
    return p


def picture(doc, filename, caption, width=6.0):
    path = ROOT / "outputs" / filename
    doc.add_picture(str(path), width=Inches(width))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_text(p, caption, font="Cambria", size=10, italic=True, color=SUBTLE)


def add_footer(doc):
    footer = doc.sections[0].footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_text(
        p,
        "Research Project Report | Lean Manufacturing for Beverage Bottling Line",
        font="Cambria",
        size=9,
        color=SUBTLE,
    )


def build():
    doc = Document()
    configure(doc)
    add_footer(doc)

    today = date(2026, 4, 25).strftime("%B %d, %Y")

    centered(doc, "PROJECT REPORT", size=13, bold=True, color=SUBTLE, space_after=10)
    centered(
        doc,
        "Simulation-Based Implementation of Lean Manufacturing Techniques",
        size=20,
        bold=True,
        color=ACCENT,
        space_after=2,
    )
    centered(
        doc,
        "for Productivity Improvement in SMEs",
        size=18,
        bold=True,
        color=ACCENT,
        space_after=3,
    )
    centered(
        doc,
        "A Data-Driven Case Study from a Beverage Bottling Line",
        size=13.5,
        bold=False,
        color=SUBTLE,
        space_after=16,
    )
    centered(doc, "Submitted by: Divash Krishnam, Ashutosh Verma, Raja", size=11.5, color=BODY, space_after=3)
    centered(doc, "Guided by: Dr. MOHD SHUAIB", size=11.5, color=BODY, space_after=3)
    centered(doc, "Department of Production and Industrial Engineering", size=11.5, color=BODY, space_after=3)
    centered(doc, f"Date: {today}", size=11.5, color=BODY, space_after=10)
    centered(
        doc,
        "This report is based on the research paper and supporting simulation analysis.",
        size=10.5,
        color=SUBTLE,
        space_after=0,
    )
    doc.add_page_break()

    heading(doc, "Abstract")
    body(
        doc,
        "Unlike studies that rely on theoretical frameworks, this report investigates how lean manufacturing priorities can be selected in an SME beverage bottling line by using real production and downtime records instead of relying purely on conceptual analysis. The dataset includes 59 production summaries, 265 hourly observations, and 1,388 downtime events recorded between July 22, 2022, and February 6, 2023. To evaluate the likely impact of lean interventions, an event-level bootstrap simulation with 5,000 replications was used together with downtime classification into micro-, minor-, and major-stoppage groups.",
    )
    body(
        doc,
        "The line produced 212,358 liters, but only 42.66% of monitored time was value-adding production time. The remaining 57.64% was consumed by waiting waste in the form of downtime. Micro-stoppages accounted for 61.31% of all events but only 13.59% of downtime minutes, whereas major stoppages represented just 6.77% of events yet consumed 66.26% of lost time. Simulation results show that 5S with visual management improves output by 2.33%, TPM improves output by 20.61%, and an integrated lean bundle improves output by 32.19% while raising simulated availability from 42.37% to 55.33%.",
    )
    subheading(doc, "Keywords")
    body(
        doc,
        "Lean manufacturing, productivity improvement, SME manufacturing, waste reduction, waiting waste, beverage bottling line, downtime analysis, total productive maintenance, 5S, bootstrap simulation.",
    )

    heading(doc, "1. Introduction")
    body(
        doc,
        "Not every problem on a bottling line costs the same amount of time. A typical line may see many short interruptions during a shift, but only a few longer stoppages can consume most of the available production time. In lean manufacturing, that non-productive interval appears as waiting waste because monitored time passes without generating output.",
    )
    body(
        doc,
        "For SMEs, this creates an important prioritization problem. Maintenance capacity, technical manpower, and improvement budgets are limited, so the first intervention must target the most damaging loss. This report addresses that problem through a real beverage bottling dataset collected from an IIoT-enabled production system and uses plant data to test which lean method is most likely to recover productive time.",
    )
    subheading(doc, "Project Objectives")
    for item in [
        "To analyze the production behavior and downtime structure of a real bottling-line system.",
        "To identify the dominant waiting-waste pattern affecting productivity.",
        "To simulate the effect of selected lean interventions on throughput, availability, and waste reduction.",
        "To recommend an implementation priority for lean techniques based on the observed loss pattern.",
    ]:
        number(doc, item)

    heading(doc, "2. Literature Review and Research Gap")
    body(
        doc,
        "The literature linked to this study can be understood through four connected discussions: general lean performance studies, food and beverage manufacturing studies, simulation-based lean evaluation studies, and downtime-centered bottling-line studies. Together, these works confirm that lean and simulation are useful, but they still leave a practical gap in deciding which intervention should come first on a real bottling line.",
    )
    body(
        doc,
        "The key research gap is that published work still rarely combines SME relevance, beverage-line context, open event data, uncertainty-aware simulation, and explicit downtime-driven prioritization in one reproducible workflow. This report addresses that gap by connecting event-level downtime records directly to lean-priority selection.",
    )
    subheading(doc, "Key Research Gap Areas")
    key_point(doc, "Lean in SMEs", "Earlier studies support lean adoption in small firms, but they provide limited plant-level ranking of lean priorities.")
    key_point(doc, "Food and beverage context", "Sector studies confirm relevance, yet they offer weak guidance on whether a bottling line should address micro-stops or major stoppages first.")
    key_point(doc, "Simulation studies", "Simulation is useful for what-if evaluation, but many prior studies use proprietary cases and are harder to reproduce.")
    key_point(doc, "Downtime studies", "Downtime-focused work diagnoses bottling-line losses well, but usually stops short of converting stoppage structure into a lean-priority sequence.")

    heading(doc, "3. Dataset and Analytical Framework")
    body(
        doc,
        "The focus of this project is a real beverage bottling line represented through a publicly available dataset published on Zenodo in January 2026. The dataset includes hourly production output records, daily production summaries, hourly operating breakdown records, and individual breakdown event records. Because the data capture both production performance and detailed stoppage events, they are well suited to a study focused on lean prioritization through observed loss patterns.",
    )
    subheading(doc, "Dataset Snapshot")
    key_point(doc, "Monitoring period", "July 22, 2022 to February 6, 2023")
    key_point(doc, "Production records", "59 production summaries and 265 hourly operating observations")
    key_point(doc, "Downtime data", "1,388 individual downtime events")
    key_point(doc, "Products observed", "3-liter and 5-liter beverage containers")
    key_point(doc, "Total recorded output", "212,358 liters from 60,888 units")
    key_point(doc, "Time base", "238.65 monitored hours with 101.81 hours of actual operation")
    subheading(doc, "Analytical Relationships Used")
    key_point(doc, "Availability", "The share of monitored time spent in actual production.")
    key_point(doc, "Downtime", "The difference between monitored time and operating time.")
    key_point(doc, "Production rate", "Liters produced per operating hour.")
    key_point(doc, "Scenario output", "Expected output after downtime is reduced under a selected lean scenario.")
    key_point(doc, "Waste reduction", "The percentage decrease in waiting waste after improvement.")

    heading(doc, "4. Methodology")
    body(
        doc,
        "This project uses a data-driven simulation methodology to evaluate lean alternatives on the basis of observed production behavior. The approach combines baseline performance analysis, downtime classification, waiting-waste measurement, and bootstrap simulation.",
    )
    subheading(doc, "Method Steps")
    for item in [
        "Extract and clean production summaries, hourly operation records, and downtime-event data.",
        "Measure baseline production, operating time, pause time, availability, and waiting-waste share.",
        "Classify downtime into micro-stoppages (<= 2 min), minor stoppages (2 to 10 min), and major stoppages (> 10 min).",
        "Estimate day-level production rates from the observed data.",
        "Run an event-level bootstrap simulation with 5,000 replications.",
        "Apply lean scenario rules and recalculate output, availability, and waste reduction.",
    ]:
        number(doc, item)
    subheading(doc, "Lean Scenarios")
    key_point(doc, "5S + Visual Management", "Assumes 15% of micro-stoppages are eliminated to represent better workplace organization and faster response to small disruptions.")
    key_point(doc, "TPM", "Assumes a 25% reduction in major stoppage duration to represent improved reliability and maintenance discipline.")
    key_point(doc, "Integrated Lean Bundle", "Combines 20% micro-stop removal, 15% minor-stop reduction, and 30% major-stop reduction to represent a broader improvement package.")
    subheading(doc, "How Lean Was Mapped to Data")
    body(
        doc,
        "Lean manufacturing was applied directly to the downtime data by linking each tool to a loss category. 5S and visual management were connected to micro-stoppages and short interruptions, TPM was connected to major stoppages and long-duration downtime, and the integrated bundle was used to model combined effects across micro, minor, and major stoppages.",
    )

    heading(doc, "5. Results and Analysis")
    subheading(doc, "Baseline Performance")
    body(
        doc,
        "The recorded production system produced 212,358 liters from 60,888 units during 238.65 monitored hours. However, operating time was only 101.81 hours, while pause time reached 137.56 hours. This means overall availability was 42.66%, and 57.64% of monitored time was consumed by waiting waste.",
    )
    key_point(doc, "Total production", "212,358 liters from 60,888 units")
    key_point(doc, "Monitored time", "238.65 hours")
    key_point(doc, "Operating time", "101.81 hours")
    key_point(doc, "Pause time", "137.56 hours")
    key_point(doc, "Overall availability", "42.66%")
    key_point(doc, "Waiting-waste share", "57.64% of monitored time")
    key_point(doc, "Observed production rates", "1,963.71 L/h for the 3-liter product and 2,918.68 L/h for the 5-liter product")
    picture(doc, "monthly_efficiency.png", "Figure 1. Monthly average operational efficiency across the monitored bottling-line period.", 6.0)
    body(
        doc,
        "Monthly efficiency varies noticeably across the study period. Lower performance appears in August 2022 and December 2022, while stronger performance is visible in January 2023. This confirms that the production line does not behave like a stable deterministic system.",
    )
    picture(doc, "downtime_distribution.png", "Figure 2. Downtime-event duration distribution. Most events are short, but a small number of long stoppages dominate loss.", 6.0)
    subheading(doc, "Downtime Structure")
    key_point(doc, "Micro-stoppages", "851 events, equal to 61.31% of all downtime events, but only 13.59% of downtime minutes")
    key_point(doc, "Minor stoppages", "443 events contributing 20.15% of downtime minutes")
    key_point(doc, "Major stoppages", "94 events, only 6.77% of all events, yet responsible for 66.26% of total lost time")
    picture(doc, "downtime_category_pareto.png", "Figure 3. Downtime Pareto by event category. Major stoppages are rare but dominant in lost time.", 5.9)
    body(
        doc,
        "A clear pattern emerges once downtime is categorized. Micro-stoppages account for 61.31% of all events but only 13.59% of downtime minutes, whereas major stoppages account for just 6.77% of events yet consume 66.26% of lost time. Event frequency alone therefore cannot identify the largest productivity loss.",
    )
    picture(doc, "pause_vs_efficiency.png", "Figure 4. Daily pause time versus production efficiency.", 5.9)
    subheading(doc, "Simulation Results")
    body(
        doc,
        "The simulation baseline produced 211,924.69 liters, which is very close to the real system's 212,358 liters. This indicates that the bootstrap simulation preserves actual plant behavior well enough for scenario comparison.",
    )
    key_point(doc, "Baseline simulation", "211,924.69 liters of output with 42.37% simulated availability")
    key_point(doc, "5S + Visual Management", "Output improves by 2.33%, waiting waste is reduced by 1.52%, and 2.10 hours of waiting time are recovered")
    key_point(doc, "TPM", "Output improves by 20.61%, waiting waste is reduced by 14.68%, and 20.21 hours of lost time are recovered")
    key_point(doc, "Integrated Lean Bundle", "Output improves by 32.19%, simulated availability rises to 55.33%, and 30.99 hours of waiting waste are recovered")
    picture(doc, "scenario_output.png", "Figure 5. Simulated production output across the baseline and lean scenarios.", 5.9)
    picture(doc, "scenario_waste_reduction.png", "Figure 6. Waiting-waste recovery across the simulated lean scenarios.", 5.9)
    body(
        doc,
        "The results show that 5S and visual management provide only a small gain because the scenario mainly targets micro-stoppages, which are frequent but relatively low in lost-time severity. TPM produces a much larger improvement because it directly targets major stoppages, which dominate downtime minutes. The integrated lean bundle performs best overall because it reduces waste across all three stoppage categories.",
    )

    heading(doc, "6. Discussion and Recommendations")
    body(
        doc,
        "The findings suggest that lean implementation should be based on the structure of operational loss rather than on the visibility or frequency of events alone. In many production systems, frequent short interruptions attract attention because they are noticeable in daily work. In this bottling line, however, the largest production losses and the largest share of waiting waste are associated with major stoppages.",
    )
    body(
        doc,
        "This explains why TPM outperforms 5S in the present study. Severe stoppages are relatively rare, but they consume the majority of lost time. By focusing first on equipment reliability and breakdown recovery, the plant can recover much more productive time than by concentrating only on workplace-organization measures.",
    )
    subheading(doc, "Recommended Order of Implementation")
    for item in [
        "First priority: TPM targeting long-duration stoppages through preventive maintenance planning, equipment checks, fast fault recovery, and reliability tracking.",
        "Second priority: 5S and visual management to reduce routine micro-disruptions, improve operator control, and standardize the workplace.",
        "Third priority: integrated lean bundle after maintenance discipline is established so that the plant can combine machine reliability with micro-stop reduction for higher gains.",
    ]:
        number(doc, item)
    body(
        doc,
        "In practical terms, the roadmap begins with TPM because major stoppages are the true drivers of lost output in this case. After severe breakdowns are better controlled, 5S and visual management become useful supporting practices for reducing routine disruptions and stabilizing day-to-day line behavior.",
    )

    heading(doc, "7. Conclusion, Limitations, and Future Scope")
    body(
        doc,
        "This report shows how lean-manufacturing priorities can be selected in a beverage bottling line by using real production and downtime data. The analysis demonstrates that the line is dominated by waiting waste because more than half of monitored time is consumed without producing output. Although micro-stoppages occur most often, they are not the main source of lost production time.",
    )
    body(
        doc,
        "The simulation confirms that the greatest improvement comes from resolving severe stoppages before focusing on the frequent but lower-impact short interruptions. For this production line, TPM is the most logical first action, while 5S and visual management become more valuable as supporting measures after major downtime is brought under better control.",
    )
    subheading(doc, "Limitations")
    for item in [
        "The lean scenarios are simulated assumptions rather than observed post-implementation outcomes.",
        "The dataset does not include detailed cause-coded downtime labels, so the mapping from downtime category to lean tool remains implicit.",
        "The analysis is based on one industrial case and should not be generalized mechanically to all production systems.",
    ]:
        bullet(doc, item)
    subheading(doc, "Future Scope")
    for item in [
        "Collect root-cause downtime data and operator-level observations.",
        "Extend the study to quality loss, changeover analysis, workforce allocation, and cost-benefit evaluation.",
        "Validate the simulated improvement levels through real shop-floor implementation.",
    ]:
        bullet(doc, item)

    heading(doc, "References")
    references = [
        "[1] I. Belekoukias, J. A. Garza-Reyes, and V. Kumar, International Journal of Production Research, 2014.",
        "[2] O. Oleghe and K. Salonitis, International Journal of Lean Six Sigma, 2019.",
        "[3] F. Yu and Z. Chen, Procedia CIRP, 2021.",
        "[4] J. Possik et al., International Journal of Computer Integrated Manufacturing, 2022.",
        "[5] S. Frecassetti et al., Quality Management Journal, 2023.",
        "[6] J. A. C. Bokhorst et al., International Journal of Production Economics, 2022.",
        "[7] G. A. David et al., IEEE Internet of Things Journal, 2024.",
        "[8] G. A. David, P. M. C. Monson, and C. Soares Junior, Industrial Production Time-Series Dataset from a Beverage Bottling Line, Zenodo, 2026.",
        "[9] J. M. Matindana and M. J. Shoshiwa, Management System Engineering, 2025.",
        "[10] M. Reslan et al., Procedia CIRP, 2025.",
        "[18] A. Belhadi et al., Journal of Industrial and Production Engineering, 2018.",
        "[19] M. Dora et al., Trends in Food Science and Technology, 2013.",
        "[20] T. Al-Hawari et al., 2010 International Conference on Computer Modeling and Simulation.",
        "[21] I. Zennaro et al., International Journal of Quality and Reliability Management, 2018.",
        "[22] G. K. Inyiama and S. A. Oke, International Journal of Quality and Reliability Management, 2020.",
        "[23] C. Cuggia-Jimenez et al., Informacion Tecnologica, 2020.",
    ]
    for ref in references:
        body(doc, ref)

    doc.save(OUTPUT)


if __name__ == "__main__":
    build()
    print(f"Created: {OUTPUT}")
