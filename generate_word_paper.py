from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parent
OUTPUT = ROOT / "final_research_paper_divash_team_word_ready.docx"


def set_page_margins(section, margin_inch=1.0):
    margin = Inches(margin_inch)
    section.top_margin = margin
    section.bottom_margin = margin
    section.left_margin = margin
    section.right_margin = margin


def set_default_font(document):
    styles = document.styles
    for style_name in ["Normal", "Title", "Heading 1", "Heading 2", "Heading 3"]:
        style = styles[style_name]
        font = style.font
        font.name = "Times New Roman"
        font.size = Pt(12 if style_name == "Normal" else 14)
        element = style.element
        rpr = element.get_or_add_rPr()
        rfonts = rpr.rFonts
        if rfonts is None:
            rfonts = OxmlElement("w:rFonts")
            rpr.append(rfonts)
        rfonts.set(qn("w:ascii"), "Times New Roman")
        rfonts.set(qn("w:hAnsi"), "Times New Roman")


def add_centered(document, text, size=12, bold=False, space_after=6):
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    p.paragraph_format.space_after = Pt(space_after)
    return p


def add_heading(document, text, level=1):
    p = document.add_paragraph()
    p.style = f"Heading {level}"
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    return p


def add_paragraph(document, text, italic=False, bold=False):
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    run.italic = italic
    run.bold = bold
    return p


def add_bullet(document, text):
    p = document.add_paragraph(style="List Bullet")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    return p


def add_number(document, text):
    p = document.add_paragraph(style="List Number")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    return p


def add_table(document, title, headers, rows):
    add_paragraph(document, title, bold=True)
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = str(value)
    document.add_paragraph()


def add_picture(document, image_path, caption, width=6.2):
    document.add_picture(str(image_path), width=Inches(width))
    last_paragraph = document.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = document.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap.add_run(caption)
    run.italic = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(11)


def build_document():
    doc = Document()
    set_default_font(doc)
    section = doc.sections[0]
    set_page_margins(section)

    add_centered(doc, "Simulation-Based Implementation of Lean Manufacturing Techniques for Productivity Improvement in SMEs:", size=16, bold=True, space_after=2)
    add_centered(doc, "A Data-Driven Case Study From a Beverage Bottling Line", size=16, bold=True, space_after=18)
    add_centered(doc, "Final Research Paper", size=14, bold=True, space_after=18)
    add_centered(doc, "Submitted by: Divash Krishnam", size=12, bold=True)
    add_centered(doc, "Team Members: Ashutosh Verma, Raja", size=12)
    add_centered(doc, "Department of Production and Industrial Engineering", size=12)
    add_centered(doc, "Guided by: Dr. MOHD SHUAIB", size=12, bold=True)
    add_centered(doc, "Mechanical Department", size=12)
    add_centered(doc, "Date: March 26, 2026", size=12)
    doc.add_page_break()

    add_heading(doc, "Abstract", 1)
    add_paragraph(
        doc,
        "This paper studies how lean-manufacturing priorities can be selected in an SME beverage bottling line by using real production and downtime records rather than a hypothetical example. The dataset contains 59 production summaries, 265 hourly observations, and 1,388 downtime events collected between July 22, 2022 and February 6, 2023. The analysis begins with throughput, availability, pause time, and waiting-waste measurement, then classifies downtime into micro-, minor-, and major-stoppage groups. An event-level bootstrap simulation is used to compare three lean scenarios: 5S with visual management, total productive maintenance (TPM), and an integrated lean bundle. The observed system produced 212,358 liters, but only 42.66% of monitored time was value-adding; the remaining 57.64% was consumed by waiting waste in the form of downtime. Simulation results show that 5S with visual management improves output by 2.33% and reduces waiting waste by 1.52%, TPM improves output by 20.61% and reduces waiting waste by 14.68%, and the integrated lean bundle improves output by 32.19%, raises simulated availability from 42.37% to 55.33%, and recovers 30.99 hours of waiting waste. The overall conclusion is that maintenance-focused action aimed at severe stoppages should be prioritized first on this line."
    )

    add_heading(doc, "Keywords", 1)
    add_paragraph(doc, "Lean manufacturing, productivity improvement, SME manufacturing, waste reduction, waiting waste, beverage bottling line, downtime analysis, total productive maintenance, 5S, bootstrap simulation.")

    add_heading(doc, "1. Introduction", 1)
    intro_paras = [
        "On a bottling line, productivity loss is not spread evenly across all operating problems. A plant may experience many short interruptions during a shift, but a much smaller number of long stoppages can consume most of the available production time. In lean terms, that lost time appears as waiting waste because monitored time passes without generating output.",
        "For SMEs, this decision is especially important. These firms usually work with limited maintenance capacity, tighter budgets, and less room for unsuccessful improvement trials. If the first intervention is chosen poorly, the plant may spend effort without recovering much productive time.",
        "The present study uses an open dataset from a beverage bottling line monitored through IIoT infrastructure. Because the dataset includes production output, operating time, pause time, and event-level downtime records, it becomes possible to connect lean actions directly to observed operating losses and to estimate how much waiting waste each intervention can remove.",
        "Rather than discussing lean only at a general level, the paper uses real plant data to test which intervention is most likely to improve throughput, raise availability, and reduce non-value-adding time on the observed line.",
    ]
    for para in intro_paras:
        add_paragraph(doc, para)
    add_paragraph(doc, "The main objectives of this study are:", bold=True)
    for item in [
        "to analyze the production behavior and downtime structure of a real bottling-line system;",
        "to identify the dominant waiting-waste pattern affecting productivity;",
        "to simulate the effect of selected lean interventions on throughput, availability, and waste reduction; and",
        "to recommend an implementation priority for lean techniques based on data-driven evidence.",
    ]:
        add_number(doc, item)

    add_heading(doc, "2. Literature Review and Research Gap", 1)
    lit_paras = [
        "The literature relevant to this study can be organized into four linked conversations rather than a single uniform debate. The first concerns the general relationship between lean manufacturing and operational performance. Belekoukias et al. [1] showed that lean methods are frequently associated with improvements in manufacturing performance, while Belhadi et al. [18] emphasized that SMEs face recurring constraints in resources, sequencing, and measurement. Together, these studies suggest that SMEs need help deciding which intervention should come first instead of treating lean as a single package.",
        "A second conversation focuses on food and beverage production, where hygiene requirements, packaging synchronization, and short stoppages complicate the transfer of lean lessons from other industries. Dora et al. [19] found that lean practice in European food-processing SMEs remained relatively immature and strongly influenced by workforce capability and organizational culture. Cuggia-Jimenez et al. [23] similarly showed through a systematic review that food-industry lean research is broad in coverage but uneven in method. Matindana and Shoshiwa [9] linked lean practices to better outcomes in food and beverage SMEs, but their sector-level model does not provide plant-level intervention ranking.",
        "A third conversation concerns simulation and model-based evaluation of lean alternatives. Oleghe and Salonitis [2] argued that lean analysis becomes stronger when process variability and system interaction are modeled explicitly. Yu and Chen [3] reached a related conclusion from an SME perspective, noting that simulation is valuable but still underused in smaller firms. Possik et al. [4] and Frecassetti et al. [5] showed that simulation can compare lean alternatives before implementation, while Al-Hawari et al. [20] used simulation in a bottle-filling system to connect downtime analysis to throughput improvement.",
        "A fourth conversation is more directly downtime-centered. Zennaro et al. [21] showed that detailed micro-downtime analysis in bottling lines can reveal a large share of inefficiency that would otherwise remain hidden. Inyiama and Oke [22] examined downtime in a process bottling plant through cause-and-effect and Weibull analysis, highlighting the value of maintenance-oriented diagnosis. Bokhorst et al. [6] and Reslan et al. [10] add a digital-manufacturing perspective by showing how smart-manufacturing and digital-twin methods can strengthen lean diagnosis.",
        "Taken together, the literature shows three recurring trade-offs. Review and survey studies provide breadth but limited plant-level decision guidance. Simulation studies provide stronger what-if comparison but often rely on proprietary industrial data. Bottling-line downtime studies provide deep operational diagnosis, yet they do not usually convert stoppage patterns into a ranked lean-priority framework. This is the gap addressed by the present study.",
    ]
    for para in lit_paras:
        add_paragraph(doc, para)
    add_table(
        doc,
        "Table 1. Literature Synthesis and Research Gap",
        ["Study", "Main Contribution", "Gap Remaining"],
        [
            ["Belhadi et al. [18]", "Broad review of lean production in SMEs and their implementation challenges.", "Strong on SME context, but not designed for plant-level prioritization of specific downtime losses."],
            ["Dora et al. [19]", "Connects lean practice to operational performance and critical success factors in food-processing SMEs.", "Valuable sector relevance, but limited guidance on event-level stoppage prioritization."],
            ["Yu and Chen [3]", "Review of factory simulation in SMEs and identification of adoption barriers.", "Supports simulation use, but does not provide a bottling-line loss-ranking workflow."],
            ["Possik et al. [4]", "Co-simulation framework for evaluating lean technique impact.", "Limited open-data reproducibility."],
            ["Zennaro et al. [21]", "Detailed micro-downtime analysis in beverage bottling lines with OEE relevance.", "Strong downtime diagnosis, but not framed as lean-intervention prioritization under uncertainty."],
            ["This study", "Open-data, downtime-driven simulation of lean alternatives.", "Provides a reproducible framework for student and SME use."],
        ],
    )

    add_heading(doc, "3. Problem Definition and Mathematical Formulation", 1)
    for para in [
        "The production line studied in this paper suffers from large variations in operating time and downtime. The central problem is to determine which lean strategy should be given priority in order to improve productivity most effectively.",
        "The key relationships used in this study are listed below in Word-ready form.",
    ]:
        add_paragraph(doc, para)
    for equation in [
        "Operational Availability: A = T_op / T_mon",
        "Downtime: T_dt = T_mon - T_op",
        "Production Rate: R = Q / T_op",
        "Scenario-Based Output: Q' = R x (T_mon - T'_dt)",
        "Downtime Transformation: T'_dt = T_res + Î£(alpha_i x t_i)",
        "Percentage Productivity Improvement: %Î”Q = ((Q' - Q) / Q) x 100",
    ]:
        add_paragraph(doc, equation, bold=True)

    add_heading(doc, "4. Dataset Description and Study Context", 1)
    for para in [
        "The dataset used in this study is the Industrial Production Time-Series Dataset from a Beverage Bottling Line, published on Zenodo on January 4, 2026 [8]. The data correspond to a real industrial bottling-line operation monitored through an IIoT architecture. The associated research article by David et al. [7] confirms the industrial and digital monitoring context of the data source.",
        "The dataset contains four important groups of information: production output aggregated by hour, daily production summaries, hourly operating breakdown records, and event-level downtime logs.",
        "The monitoring period covers July 22, 2022 to February 6, 2023. The dataset includes two product sizes, 3-liter and 5-liter containers, and captures real operating instability over time. Because the data are open and reusable, they are highly suitable for an academic study requiring transparency and reproducibility.",
    ]:
        add_paragraph(doc, para)
    add_table(
        doc,
        "Table 2. Dataset Characteristics",
        ["Attribute", "Value"],
        [
            ["Monitoring period", "July 22, 2022 to February 6, 2023"],
            ["Production summary records", "59"],
            ["Calendar production days", "56"],
            ["Hourly operating observations", "265"],
            ["Downtime events", "1,388"],
            ["Observed product sizes", "3 L and 5 L"],
            ["Total liters produced", "212,358"],
            ["Total units produced", "60,888"],
            ["Total monitored time", "238.65 h"],
            ["Total operating time", "101.81 h"],
        ],
    )

    add_heading(doc, "5. Research Methodology", 1)
    methodology = {
        "Step 1: Data Extraction and Cleaning": "The Excel dataset was first read and organized into production summaries, hourly operation records, and downtime-event data. Dates, durations, and production quantities were standardized. Downtime durations were converted to minutes so that event-level categorization and simulation could be performed consistently.",
        "Step 2: Baseline Performance Analysis": "Baseline performance was measured using liters produced, production units, monitored time, operating time, pause time, daily efficiency, hourly efficiency, and the share of monitored time consumed by waiting waste. This step was necessary to understand the real operating condition of the system before any lean scenario was introduced.",
        "Step 3: Downtime Categorization": "Downtime events were classified into three categories: micro-stoppages (t <= 2 min), minor stoppages (2 < t <= 10 min), and major stoppages (t > 10 min). This classification distinguishes frequent short interruptions from serious breakdown-type events.",
        "Step 4: Production-Rate Estimation": "For each observed production day, production rate was estimated from liters produced divided by effective operating time. This creates a realistic day-level performance profile that preserves operational variability instead of assuming a constant theoretical output rate.",
        "Step 5: Bootstrap Simulation": "An event-level bootstrap simulation with 5,000 replications was performed. Each replication resampled observed day profiles with replacement. This allows the model to preserve the stochastic behavior of monitored time, productivity, and downtime structure. Event durations were normalized to the recorded daily pause time so that the simulated baseline remained aligned with the observed plant record.",
        "Step 6: Scenario Design": "Three lean scenarios were designed and evaluated: 5S + Visual Management with 15% of micro-stoppages eliminated; TPM with 25% reduction in the duration of major stoppages; and an Integrated Lean Bundle with 20% of micro-stoppages eliminated, 15% reduction in minor stoppages, and 30% reduction in major stoppages. These values are scenario assumptions informed by lean implementation logic and literature.",
    }
    for title, text in methodology.items():
        add_paragraph(doc, title, bold=True)
        add_paragraph(doc, text)

    add_paragraph(doc, "Algorithm and Logic Used in the Model", bold=True)
    add_paragraph(doc, "The main algorithm used in this research is a data-driven bootstrap simulation. First, the real production and downtime records are converted into daily operating profiles. Second, these real daily profiles are resampled repeatedly so that actual plant variability is preserved. Third, lean-improvement rules are applied to the sampled downtime events. Finally, output, pause time, efficiency, and waiting-waste reduction are recalculated.")
    add_paragraph(doc, "In simple words, the algorithm repeatedly asks: if we reduce the type of downtime targeted by a lean tool, how much production can be recovered?")
    for item in [
        "Read the production sheets and downtime-event sheet.",
        "Convert downtime durations into minutes and aggregate them by date.",
        "Calculate daily monitored time, operating time, pause time, and production rate.",
        "Create baseline operating profiles for each observed day.",
        "Sample those observed day profiles with replacement across many replications.",
        "Apply lean rules to the sampled downtime events.",
        "Recalculate pause time, output, and efficiency after each scenario.",
        "Compare the simulated results with the baseline system.",
    ]:
        add_number(doc, item)

    add_paragraph(doc, "How Lean Manufacturing Was Applied to the Data", bold=True)
    add_paragraph(doc, "Lean manufacturing was used directly on the observed downtime data instead of being discussed only as a theory. In this study, downtime was treated as waiting waste because monitored production time was consumed without generating output, and each lean tool was linked to a specific category of loss in the dataset.")
    add_table(
        doc,
        "Table 2A. How Lean Tools Were Mapped to the Data",
        ["Lean Tool", "Data Feature Used", "Logic Used in the Study"],
        [
            ["5S + Visual Management", "Micro-stoppages", "Used to represent the reduction of small avoidable interruptions caused by poor workplace organization and visual control."],
            ["TPM", "Major stoppages", "Used to represent reduced machine breakdown duration and faster recovery through better maintenance."],
            ["Integrated Lean Bundle", "Micro, minor, and major stoppages", "Used to represent a combined improvement in workplace discipline and machine reliability."],
        ],
    )

    add_heading(doc, "6. Results and Visual Analysis", 1)
    add_paragraph(doc, "The observed production system produced 212,358 liters from 60,888 units during 238.65 monitored hours. Effective operating time was only 101.81 hours, while pause time reached 137.56 hours. This corresponds to an overall observed availability of 42.66%, indicating that downtime is the main performance constraint. In lean terms, 57.64% of monitored time was consumed by waiting waste.")
    add_paragraph(doc, "At the product level, the average production rate was 1,963.71 L/h for the 3-liter product and 2,918.68 L/h for the 5-liter product. The 5-liter product therefore showed a higher effective throughput during operating time.")
    add_table(
        doc,
        "Table 3. Baseline Performance Indicators",
        ["Metric", "Value"],
        [
            ["Production days", "56"],
            ["Production summary records", "59"],
            ["Hourly observations", "265"],
            ["Downtime events", "1,388"],
            ["Total production", "212,358 L"],
            ["Total units produced", "60,888"],
            ["Monitored time", "238.65 h"],
            ["Operating time", "101.81 h"],
            ["Pause time", "137.56 h"],
            ["Overall availability", "42.66%"],
            ["Waiting-waste share", "57.64%"],
        ],
    )

    images = [
        ("monthly_efficiency.png", "Figure 1. Monthly average operational efficiency from the real bottling-line dataset.", 6.2),
        ("downtime_distribution.png", "Figure 2. Distribution of downtime-event duration, truncated at 60 minutes for readability.", 6.2),
    ]
    for filename, caption, width in images:
        add_picture(doc, ROOT / "outputs" / filename, caption, width)
        doc.add_paragraph()

    add_table(
        doc,
        "Table 4. Downtime Category Structure",
        ["Category", "Event Count", "Share of Downtime Minutes"],
        [
            ["Micro (<= 2 min)", "851", "13.59%"],
            ["Minor (2-10 min)", "443", "20.15%"],
            ["Major (> 10 min)", "94", "66.26%"],
        ],
    )

    add_picture(doc, ROOT / "outputs" / "downtime_category_pareto.png", "Figure 3. Downtime Pareto by event category. Major stoppages are rare but dominant in lost time.", 6.2)
    doc.add_paragraph()
    add_paragraph(doc, "The downtime Pareto view provides one of the most important insights of the paper. Micro-stoppages account for 61.31% of all events but only 13.59% of total downtime minutes, whereas major stoppages represent only 6.77% of events but consume 66.26% of lost time. This means that event frequency alone does not identify the highest productivity loss or the largest source of waiting waste. Lean prioritization must be based on lost-time severity.")

    add_picture(doc, ROOT / "outputs" / "pause_vs_efficiency.png", "Figure 4. Daily pause time versus operational efficiency.", 6.2)
    doc.add_paragraph()
    add_paragraph(doc, "The scatter plot shows a clear downward trend. This provides further evidence that reducing downtime is the most direct productivity-improvement lever in the observed system.")

    add_picture(doc, ROOT / "outputs" / "product_rate_comparison.png", "Figure 5. Mean and median production rate by product size.", 5.8)
    doc.add_paragraph()
    add_paragraph(doc, "The 5-liter product has both higher mean and higher median production rate than the 3-liter product. This suggests that product mix can also influence line performance and may become an additional decision variable in future scheduling or line-balancing studies.")

    add_paragraph(doc, "The simulation baseline produced 211,924.69 liters on average across the resampled production horizon. This closely matches the observed output of the real system, indicating that the simulation reasonably preserves actual operating behavior. In the baseline simulation, waiting waste averaged 137.67 hours.")
    add_table(
        doc,
        "Table 5. Simulation Results for Lean Scenarios",
        ["Scenario", "Mean Output (L)", "Gain (%)", "Efficiency (%)", "Eff. Gain (pp)"],
        [
            ["Baseline", "211,924.69", "0.00", "42.37", "0.00"],
            ["5S + Visual Management", "216,866.69", "2.33", "43.25", "0.88"],
            ["TPM", "255,604.10", "20.61", "50.81", "8.44"],
            ["Integrated Lean Bundle", "280,142.22", "32.19", "55.33", "12.96"],
        ],
    )
    add_table(
        doc,
        "Table 5A. Waste-Reduction Results for Lean Scenarios",
        ["Scenario", "Waiting Waste (h)", "Waste Recovered (h)", "Waste Reduction (%)"],
        [
            ["Baseline", "137.67", "0.00", "0.00"],
            ["5S + Visual Management", "135.57", "2.10", "1.52"],
            ["TPM", "117.46", "20.21", "14.68"],
            ["Integrated Lean Bundle", "106.68", "30.99", "22.51"],
        ],
    )
    add_table(
        doc,
        "Table 6. Simulation Uncertainty Range",
        ["Scenario", "P05 Output (L)", "P95 Output (L)"],
        [
            ["Baseline", "188,572.81", "236,799.79"],
            ["5S + Visual Management", "192,906.87", "242,100.46"],
            ["TPM", "226,295.58", "287,389.24"],
            ["Integrated Lean Bundle", "247,828.10", "314,804.33"],
        ],
    )
    add_picture(doc, ROOT / "outputs" / "scenario_output.png", "Figure 6. Mean simulated production output across the baseline and lean scenarios.", 6.2)
    doc.add_paragraph()
    add_picture(doc, ROOT / "outputs" / "scenario_waste_reduction.png", "Figure 7. Recovered waiting waste and waste-reduction percentage across the lean scenarios.", 6.2)
    doc.add_paragraph()
    add_paragraph(doc, "The simulation results show that 5S and visual management provide only a modest output gain of 2.33% and recover 2.10 hours of waiting waste. This is because the scenario mainly targets micro-stoppages, which are frequent but contribute relatively little to total lost time. By contrast, TPM produces a much larger gain of 20.61% and recovers 20.21 hours of waiting waste because it directly targets major stoppages, which dominate downtime minutes. The integrated lean bundle produces the highest gain, improving output by 32.19%, raising simulated availability from 42.37% to 55.33%, and reducing waiting waste by 22.51%.")

    add_heading(doc, "7. Discussion", 1)
    for para in [
        "The results strongly suggest that lean implementation should be based on the structure of operational loss rather than on the visibility or frequency of events alone. In many production systems, micro-stoppages attract attention because they happen often and disrupt workflow visibly. However, as demonstrated in this study, those frequent interruptions may not be the dominant cause of lost production.",
        "For the bottling line studied here, major stoppages are the real productivity driver. Although they are relatively rare, they consume most of the downtime minutes and therefore most of the waiting waste. This explains why TPM, which focuses on equipment reliability and breakdown reduction, produces far more improvement than 5S alone. The result is fully consistent with the broader lean-simulation literature, which emphasizes context-specific tool selection [4], [5].",
        "Another important implication is the role of digital manufacturing data in strengthening lean analysis. Instead of relying only on static value stream mapping or qualitative observation, this study uses event-level downtime records and simulation to quantify both output recovery and waste reduction. This is aligned with the direction suggested by smart manufacturing and digital twin research [6], [10]. The practical value is clear: real data allow the analyst to identify high-leverage interventions rather than applying lean tools generically.",
        "For final-year engineering research, this approach is especially useful because it combines core industrial engineering ideas: data analysis, productivity evaluation, lean manufacturing, operations research thinking, and simulation-based decision support. It also produces a study that is more rigorous than a purely descriptive case report.",
    ]:
        add_paragraph(doc, para)

    add_heading(doc, "8. Implementation Recommendations", 1)
    add_paragraph(doc, "Based on the results, the recommended order of lean implementation is:")
    for item in [
        "First priority: TPM. The plant should focus on reducing long-duration stoppages through maintenance planning, preventive checks, fast fault recovery, and equipment reliability tracking.",
        "Second priority: 5S and visual management. These tools should be used to reduce small disruptions, improve operator control, and standardize the workplace.",
        "Third priority: integrated lean bundle. Once maintenance discipline is improved, the plant can combine TPM with micro-stop reduction and standard work for broader gains.",
    ]:
        add_number(doc, item)
    add_table(
        doc,
        "Table 7. Recommended Lean Implementation Roadmap",
        ["Priority", "Lean Focus", "Expected Impact"],
        [
            ["1", "TPM and maintenance discipline", "Highest reduction in lost-time severity and waiting waste by targeting major stoppages."],
            ["2", "5S and visual management", "Better workplace organization and reduction in frequent micro-disruptions."],
            ["3", "Integrated lean bundle", "Combined gains from reliability, stabilization, and standardized work."],
        ],
    )

    add_heading(doc, "9. Conclusion", 1)
    for para in [
        "This paper presented a simulation-based lean study using a real industrial beverage bottling dataset. The study showed that the system suffers from low availability, high downtime, and a large share of waiting waste across the monitored period. Although micro-stoppages are the most frequent events, major stoppages dominate lost production time.",
        "The simulation results demonstrated that 5S and visual management improve output by 2.33% and reduce waiting waste by 1.52%, TPM improves output by 20.61% and reduces waiting waste by 14.68%, and an integrated lean bundle improves output by 32.19% while raising simulated availability from 42.37% to 55.33% and reducing waiting waste by 22.51%.",
        "The main managerial conclusion is that maintenance-led reduction of major stoppages should be prioritized first. The main academic contribution is the development of a reproducible, open-data framework for simulation-based lean research that evaluates throughput, availability, and waste reduction together in an SME context.",
    ]:
        add_paragraph(doc, para)

    add_heading(doc, "10. Limitations and Future Scope", 1)
    for para in [
        "This study has three major limitations. First, the lean scenarios are simulated assumptions rather than observed post-implementation outcomes. Second, the dataset does not contain detailed cause-coded downtime labels, so the mapping from downtime category to lean tool remains inferential. Third, the analysis is based on one industrial case, so the results should not be generalized mechanically to all production systems.",
        "Future work can improve the present study by collecting root-cause-coded downtime data, operator-level observations, quality-loss information, and real post-intervention performance records. Future studies may also extend the simulation to include changeover analysis, workforce allocation, line balancing, and cost-benefit evaluation.",
    ]:
        add_paragraph(doc, para)

    add_heading(doc, "Acknowledgement", 1)
    add_paragraph(doc, "The authors sincerely acknowledge the guidance and support of Dr. MOHD SHUAIB from the Mechanical Department. His academic guidance was valuable in shaping the direction and structure of this work.")

    add_heading(doc, "References", 1)
    references = [
        "[1] I. Belekoukias, J. A. Garza-Reyes, and V. Kumar, The impact of lean methods and tools on the operational performance of manufacturing organisations, International Journal of Production Research, vol. 52, no. 18, pp. 5346-5366, 2014. doi: 10.1080/00207543.2014.903348.",
        "[2] O. Oleghe and K. Salonitis, Hybrid simulation modelling of the human-production process interface in lean manufacturing systems, International Journal of Lean Six Sigma, vol. 10, no. 2, pp. 665-690, 2019. doi: 10.1108/IJLSS-01-2018-0004.",
        "[3] F. Yu and Z. Chen, Tools, application areas and challenges of factory simulation in Small and Medium-Sized Enterprises - A Review, Procedia CIRP, vol. 104, pp. 399-404, 2021. doi: 10.1016/j.procir.2021.11.067.",
        "[4] J. Possik, A. Zouggar-Amrani, B. Vallespir, and G. Zacharewicz, Lean techniques impact evaluation methodology based on a co-simulation framework for manufacturing systems, International Journal of Computer Integrated Manufacturing, vol. 35, no. 1, pp. 91-111, 2022. doi: 10.1080/0951192X.2021.1972468.",
        "[5] S. Frecassetti, B. Kassem, K. Kundu, M. Ferrazzi, and A. Portioli-Staudacher, Introducing Lean practices through simulation: A case study in an Italian SME, Quality Management Journal, vol. 30, no. 2, pp. 90-104, 2023. doi: 10.1080/10686967.2023.2171326.",
        "[6] J. A. C. Bokhorst, W. Knol, J. Slomp, and T. Bortolotti, Assessing to what extent smart manufacturing builds on lean principles, International Journal of Production Economics, vol. 253, art. no. 108599, 2022. doi: 10.1016/j.ijpe.2022.108599.",
        "[7] G. A. David, P. M. C. Monson, C. Soares Junior, P. O. Conceicao Junior, P. R. Aguiar, and A. Simeone, IoT-Driven Deep Learning for Enhanced Industrial Production Forecasting, IEEE Internet of Things Journal, vol. 11, no. 23, pp. 38486-38495, 2024. doi: 10.1109/JIOT.2024.3447579.",
        "[8] G. A. David, P. M. C. Monson, and C. Soares Junior, Industrial Production Time-Series Dataset from a Beverage Bottling Line. Zenodo, Jan. 4, 2026. doi: 10.5281/zenodo.18146866.",
        "[9] J. M. Matindana and M. J. Shoshiwa, Lean manufacturing implementation in food and beverage SMEs in Tanzania: using structural equation modelling (SEM), Management System Engineering, vol. 4, no. 1, pp. 1-14, 2025. doi: 10.1007/s44176-025-00036-3.",
        "[10] M. Reslan, M. J. Triebe, R. Venketesh, and A. J. Hartwell, Automation of Value Stream Mapping: A Case Study on Enhancing Lean Manufacturing Tools Through Digital Twins, Procedia CIRP, vol. 134, pp. 455-460, 2025. doi: 10.1016/j.procir.2025.02.159.",
        "[11] C. Unal and S. Bilget, Examination of lean manufacturing systems by simulation technique in apparel industry, The Journal of The Textile Institute, vol. 112, no. 3, pp. 377-387, 2021. doi: 10.1080/00405000.2020.1756104.",
        "[18] A. Belhadi, Y. B. M. Sha'ri, F. E. Touriki, and S. El Fezazi, Lean production in SMEs: literature review and reflection on future challenges, Journal of Industrial and Production Engineering, vol. 35, no. 6, pp. 368-382, 2018. doi: 10.1080/21681015.2018.1508081.",
        "[19] M. Dora, M. Kumar, D. Van Goubergen, A. Molnar, and X. Gellynck, Operational performance and critical success factors of lean manufacturing in European food processing SMEs, Trends in Food Science and Technology, vol. 31, no. 2, pp. 156-164, 2013. doi: 10.1016/j.tifs.2013.03.002.",
        "[20] T. Al-Hawari, F. Aqlan, M. Al-Buhaisi, and Z. Al-Faqeer, Simulation-based analysis and productivity improvement of a fully automatic bottle-filling production system: A practical case study, in 2010 International Conference on Computer Modeling and Simulation, vol. 4, pp. 195-199, 2010. doi: 10.1109/ICCMS.2010.212.",
        "[21] I. Zennaro, D. Battini, F. Sgarbossa, A. Persona, and R. De Marchi, Micro downtime: Data collection, analysis and impact on OEE in bottling lines: The San Benedetto case study, International Journal of Quality and Reliability Management, vol. 35, no. 4, pp. 965-995, 2018. doi: 10.1108/IJQRM-11-2016-0202.",
        "[22] G. K. Inyiama and S. A. Oke, Maintenance downtime evaluation in a process bottling plant, International Journal of Quality and Reliability Management, vol. 38, no. 1, pp. 229-248, 2020. doi: 10.1108/IJQRM-12-2018-0340.",
        "[23] C. Cuggia-Jimenez, E. Orozco-Acosta, and D. Mendoza-Galvis, Lean manufacturing: a systematic review in the food industry, Informacion Tecnologica, vol. 31, no. 5, pp. 163-172, 2020. doi: 10.4067/S0718-07642020000500163.",
    ]
    for ref in references:
        add_paragraph(doc, ref)

    doc.save(OUTPUT)


if __name__ == "__main__":
    build_document()
    print(f"Created: {OUTPUT}")

